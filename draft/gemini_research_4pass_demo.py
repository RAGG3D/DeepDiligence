#!/usr/bin/env python3
"""
gemini_research_4pass_demo.py -- Demo: one search → four independent thinking passes

Architecture:
  Phase 1: ONE Deep Research call (web search + data collection) → raw data report
  Phase 2: FOUR separate Gemini 2.5 Flash calls, each receiving the raw data as context:
           - Base case analysis
           - Bull case analysis
           - Bear case analysis
           - Catalyst analysis

This separates the expensive web search from the reasoning, allowing each scenario
to get a fully independent deep-thinking pass with the SAME information pool.

Usage:
    python gemini_research_4pass_demo.py --ticker CMPX --company-name "Compass Therapeutics" --drug CTX-009

    # Skip Phase 1 (reuse existing data report):
    python gemini_research_4pass_demo.py --ticker CMPX --company-name "Compass Therapeutics" --drug CTX-009 \
        --data-report /path/to/existing_data_report.md
"""

import argparse
import json
import logging
import os
import re
import sys
import time
import warnings
from datetime import datetime
from pathlib import Path
from typing import Optional

try:
    from google import genai
    from google.genai import types as genai_types
except ImportError:
    print("ERROR: google-genai not installed. Run: pip install google-genai")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)


# ══════════════════════════════════════════════════════════════════════════════
#  PHASE 1 PROMPT: Data Collection (NO predictions, just gather ALL data)
# ══════════════════════════════════════════════════════════════════════════════

DATA_COLLECTION_PROMPT = """
You are a senior biotech equity research data analyst. Your ONLY job is to COLLECT and ORGANIZE data.
Do NOT make predictions, do NOT project market shares, do NOT give opinions.
Just find and present ALL available data in structured tables.

## DRUG: {drug_name}
## COMPANY: {company_name} (Ticker: {ticker})

## Drug-Specific Trial Data (from ClinicalTrials.gov)
{drug_trials_text}

---

## YOUR TASK: Exhaustive Data Collection

Use Google Search to find ALL relevant data for the following sections. Present data in tables.
Cite sources for EVERY number. If a data point is not publicly available, write "NR" (not reported).

### Section 1: Drug Profile
- Drug name, chemical class (small molecule / bispecific antibody / ADC / etc.)
- Target(s) and mechanism of action
- ALL known indications with source (NCT#, 10-K page, press release)
- Current clinical stage per indication
- Route of administration, dosing schedule

### Section 2: Clinical Efficacy Data (EXHAUSTIVE)
**For EVERY data readout/presentation of {drug_name} in EVERY indication**, create a row:

| Date | Indication | NCT# | Phase | Line | N | ORR (%) | CR (%) | DCR (%) | mPFS (mo) | mOS (mo) | ≥G3 AEs (%) | Conference/Source | Stock Δ (-1d to +1d) |

Include ALL known readouts. Omitting any publicly available data point is a failure.

### Section 3: Competitive Landscape — Marketed Drugs
**For EACH indication where {drug_name} has a trial**, list ALL marketed competitors:

| Drug | Company | Indication | Line | 2024 Revenue ($M) | ORR (%) | CR (%) | DCR (%) | mPFS (mo) | mOS (mo) | ≥G3 AEs (%) | Route | Approval Year | Source |

### Section 4: Competitive Landscape — Clinical-Stage Drugs
**For EACH indication**, list ALL clinical-stage competitors in the SAME line:

| Drug | Company | Indication | Line | Phase | NCT# | N | ORR (%) | CR (%) | mPFS (mo) | mOS (mo) | ≥G3 AEs (%) | Est. Completion | Source |

### Section 5: TAM Data Per Indication
For each indication+line where {drug_name} has a trial:

| Indication | Line | Biomarker | 2024 Patients | CAGR (%) | 2030 Patients | 2038 Patients | Source |

### Section 6: Comparable Drug Launch Histories
For understanding timeline and market share ramp-up, find data on 5-10 comparable drugs
(same molecule type, similar indications, similar company size):

| Drug | Company | Target | Indication | Phase 1 Start | Phase 2 Start | Phase 3 Start | BLA Filed | Approved | Company Size at Approval | Peak MS (%) | Years to Peak | Source |

### Section 7: Company Pipeline Context
- Other drugs in {company_name}'s pipeline (brief list)
- Manufacturing capabilities
- Partnership/collaboration status for {drug_name}
- Cash runway / latest financing
- Key management relevant to {drug_name} development

### Section 8: Upcoming Catalysts
| # | Event | Indication | Expected Date | Type (interim/primary/BLA/approval) | Confidence (H/M/L) | Source |

### Section 9: Historical Stock Price Reactions to Clinical Data
| Date | Event | Stock Price Before | Stock Price After | Change (%) | Source |

---

IMPORTANT: This is DATA COLLECTION ONLY. Present raw data in tables. Do NOT predict, do NOT analyze, do NOT project market shares. The analysis will be done separately.

Begin comprehensive data collection now. Use Google Search extensively.
"""


# ══════════════════════════════════════════════════════════════════════════════
#  PHASE 2 PROMPTS: Scenario-Specific Reasoning (each gets data as context)
# ══════════════════════════════════════════════════════════════════════════════

_COMMON_RULES = """
## STRICT RULES
- Every market share % and timeline year MUST be backed by data from the research below.
- NEVER use "illustrative" or "hypothetical". Cite specific comparables.
- Line-matched comparison ONLY: 2L vs 2L, 3L vs 3L. NEVER compare 2L vs 1L.
- List ALL relevant comparables (not just one), apply majority+proximity rule.
- TAM maturity: drugs reach peak share by Year 8 post-launch; 6% post-maturity growth.
- Tiers: Best-In-Class (Tagrisso curve) / Tier One (Alecensa curve) / Average.
"""

BASE_REASONING_PROMPT = """
You are a senior biotech equity research analyst writing a BASE CASE analysis.

## DRUG: {drug_name} ({company_name}, Ticker: {ticker})

{common_rules}

## BASE CASE ASSUMPTIONS
- Clinical outcomes match the drug's OWN data (median expectation)
- Standard regulatory timeline (prioritize comparables of same company size)
- Market penetration in line with comparable drugs at same differentiation level
- All competitors launch on their projected timelines
- No unexpected safety signals, no label expansion beyond current trials

## COLLECTED DATA (from comprehensive web search — use this as your sole information source)

{data_report}

---

## REQUIRED OUTPUT STRUCTURE

### Chapter 1: Drug Overview
- Drug name, target, mechanism
- ALL indications with NCT# source
- Key differentiation vs existing therapies

### Chapter 2: Multi-Indication Forecasting Strategy
1. List every indication: active trial? analyst discussion? data disclosed?
2. Classify: CLEAR (separate forecast) vs UNCLEAR (may combine)
3. Decision: Separate All / Primary+Secondary / All Combined — with justification
4. If combining: use specific cancer names (e.g., "HNSCC/NSCLC/OV Combined"), NEVER "Other Solid Tumors"

### Chapter 3: Per-Indication Analysis
For EACH indication (or combined group):

#### 3.X.1 TAM
Patient count, growth rate, 2030/2038 projections (from Section 5 data)

#### 3.X.2 Efficacy Data Summary
Key data points from Section 2 data, quantified comparison vs competitors

#### 3.X.3 Differentiation Assessment
- Quantified: ORR X% vs Y%, mPFS X vs Y mo
- Safety: ≥G3 AE rate comparison
- Assessment: Best-in-class / Above-average / Average / Below-average

#### 3.X.4 Market Share Projection — BASE Case (2024-2038)
| Year | Market Share | Reasoning |
| 2024 | 0% | Pre-approval |
| ... | ... | ... |
| 20XX | Peak X% | Peak based on [comparable: Drug A achieved Y% with similar ORR advantage] |
| 2038 | X% | ... |

**Peak reasoning MUST include**: ALL comparable drugs, why higher/lower/equal, multiplier justification.

### Chapter 4: Stage Timeline
| Stage | Year | Source / Reasoning |
Timeline comparables MUST include ALL similar drugs with duration data.

### Chapter 5: Source Summary
All NCT#, 10-K, press release, analyst report citations.
"""

BULL_DATA_PROMPT = """
You are a senior biotech equity research DATA COLLECTOR. Your ONLY job is to COLLECT and ORGANIZE data.
Do NOT make predictions, do NOT project market shares, do NOT give opinions or analysis.
Just find and present ALL available data in structured tables.

## DRUG: {drug_name} ({company_name}, Ticker: {ticker})

## OBJECTIVE: Collect data needed for a BULL CASE valuation
Focus on: best-performing comparables, accelerated regulatory precedents, competitor failures,
label expansion precedents, and best-case launch trajectories.

## COLLECTED DATA (from Phase 1 comprehensive web search)
{data_report}

---

## YOUR TASK: Search for and present the following data tables.
Use Google Search extensively. Cite sources for EVERY number. If data is not publicly available, write "/".

### Table 1: Best-Performing Comparable Drug Readouts (Peer View Format)

For EACH indication where {drug_name} has a trial, find the TOP 3-5 drugs with the BEST efficacy.
Present as a TRANSPOSED table (one column per drug readout):

| Field | Drug A (IND) | Drug B (IND) | Drug C (IND) | Drug D (IND) | ... |
|---|---|---|---|---|---|
| Drug Name | | | | | |
| Company | | | | | |
| Ticker | | | | | |
| Innovation (BIC/FIC/Me-too) | | | | | |
| Target | | | | | |
| Result (Approved/Phase X) | | | | | |
| NCT# | | | | | |
| Treatment Line | | | | | |
| Phase | | | | | |
| Stage (1-5) | | | | | |
| Data Date | | | | | |
| N (patients) | | | | | |
| ORR (%) | | | | | |
| BICR ORR (%) | | | | | |
| CR (%) | | | | | |
| PR (%) | | | | | |
| DCR (%) | | | | | |
| mPFS (months) | | | | | |
| mRPFS (months) | | | | | |
| mOS (months) | | | | | |
| 6 Mo PFS Rate (%) | | | | | |
| 12 Mo PFS Rate (%) | | | | | |
| 24 Mo PFS Rate (%) | | | | | |
| 18 Mo OS Rate (%) | | | | | |
| 24 Mo OS Rate (%) | | | | | |
| mDFS (months) | | | | | |
| Median Follow-Up (months) | | | | | |
| ≥G3 SAE / Patients (%) | | | | | |
| ≥G3 Clinical AE (%) | | | | | |
| Route (IV/Oral/SC) | | | | | |
| Latest Annual Sale ($M) | | | | | |
| 1st Year Sale ($M) | | | | | |
| Market Reaction −1d→+1d (%) | | | | | |

Use the MOST DATA-RICH readout to define all rows. Other readouts use "/" for unreported fields.
Include {drug_name} itself as the FIRST column for direct comparison.

### Table 2: Breakthrough/Accelerated Approval Precedents

Find all drugs in the SAME indications or drug class that received expedited regulatory pathways:

| Drug | Indication | Pathway (BT/AA/Priority/Fast Track) | Time: IND→Approval (months) | Phase 2→Approval (months) | Pivotal Trial N | Company Size at Filing | Year Approved | Source |
|---|---|---|---|---|---|---|---|---|

### Table 3: Notable Competitor Failures/Delays (Same Indications)

Find drugs in the same indications that FAILED or were DELAYED:

| Drug | Company | Indication | Phase Failed | Failure Reason (efficacy/safety/enrollment) | Date | Trial NCT# | Detail (ORR/safety data if available) | Source |
|---|---|---|---|---|---|---|---|---|

### Table 4: Label Expansion Precedents

Find drugs with the same target or MOA that expanded into additional indications after initial approval:

| Drug | Target/MOA | Initial Indication | Expanded Indication | Time: Initial→Expansion (months) | Revenue Impact ($M, before→after) | Source |
|---|---|---|---|---|---|---|

### Table 5: Best-Case Launch Trajectories

Find the FASTEST revenue ramp-ups among comparable drugs (same drug class, similar indications):

| Drug | Company | Indication | Approval Year | Year 1 Revenue ($M) | Year 2 Revenue ($M) | Year 3 Revenue ($M) | Peak Revenue ($M) | Years to Peak | Peak Market Share (%) | Company Size | Source |
|---|---|---|---|---|---|---|---|---|---|---|---|

---

IMPORTANT: This is DATA COLLECTION ONLY. Present raw data in tables. Do NOT predict, do NOT analyze, do NOT project market shares, do NOT give opinions. The analysis will be done separately by a human analyst.
"""

BEAR_DATA_PROMPT = """
You are a senior biotech equity research DATA COLLECTOR. Your ONLY job is to COLLECT and ORGANIZE data.
Do NOT make predictions, do NOT project market shares, do NOT give opinions or analysis.
Just find and present ALL available data in structured tables.

## DRUG: {drug_name} ({company_name}, Ticker: {ticker})

## OBJECTIVE: Collect data needed for a BEAR CASE valuation
Focus on: full competitive landscape, safety concerns, disappointing launches,
small company challenges, and payer/reimbursement barriers.

## COLLECTED DATA (from Phase 1 comprehensive web search)
{data_report}

---

## YOUR TASK: Search for and present the following data tables.
Use Google Search extensively. Cite sources for EVERY number. If data is not publicly available, write "/".

### Table 1: Full Competitive Landscape (Peer View Format — ALL Drugs)

For EACH indication where {drug_name} has a trial, list ALL marketed AND late-stage clinical drugs.
Present as a TRANSPOSED table (one column per drug readout):

| Field | Drug A (IND) | Drug B (IND) | Drug C (IND) | Drug D (IND) | ... |
|---|---|---|---|---|---|
| Drug Name | | | | | |
| Company | | | | | |
| Ticker | | | | | |
| Innovation (BIC/FIC/Me-too) | | | | | |
| Target | | | | | |
| Result (Approved/Phase X) | | | | | |
| NCT# | | | | | |
| Treatment Line | | | | | |
| Phase | | | | | |
| Stage (1-5) | | | | | |
| Data Date | | | | | |
| N (patients) | | | | | |
| ORR (%) | | | | | |
| BICR ORR (%) | | | | | |
| CR (%) | | | | | |
| PR (%) | | | | | |
| DCR (%) | | | | | |
| mPFS (months) | | | | | |
| mRPFS (months) | | | | | |
| mOS (months) | | | | | |
| 6 Mo PFS Rate (%) | | | | | |
| 12 Mo PFS Rate (%) | | | | | |
| 24 Mo PFS Rate (%) | | | | | |
| 18 Mo OS Rate (%) | | | | | |
| 24 Mo OS Rate (%) | | | | | |
| mDFS (months) | | | | | |
| Median Follow-Up (months) | | | | | |
| ≥G3 SAE / Patients (%) | | | | | |
| ≥G3 Clinical AE (%) | | | | | |
| Route (IV/Oral/SC) | | | | | |
| Latest Annual Sale ($M) | | | | | |
| 1st Year Sale ($M) | | | | | |
| Market Reaction −1d→+1d (%) | | | | | |

Use the MOST DATA-RICH readout to define all rows. Other readouts use "/" for unreported fields.
Include {drug_name} itself as the FIRST column. Completeness is critical — omitting a competitor is a failure.

### Table 2: Detailed Safety Profile Comparison

For {drug_name} AND each major competitor, list ALL reported Grade 3+ adverse events:

| Adverse Event | {drug_name} (%) | Competitor A (%) | Competitor B (%) | Competitor C (%) | ... |
|---|---|---|---|---|---|
| Neutropenia | | | | | |
| Thrombocytopenia | | | | | |
| Anemia | | | | | |
| Hepatotoxicity (ALT/AST elevation) | | | | | |
| Diarrhea | | | | | |
| Fatigue | | | | | |
| Nausea/Vomiting | | | | | |
| Infusion-related reaction | | | | | |
| Pneumonitis/ILD | | | | | |
| Rash/Dermatitis | | | | | |
| Neuropathy | | | | | |
| Cardiac events | | | | | |
| Treatment discontinuation rate (%) | | | | | |
| Treatment-related deaths | | | | | |
| Source (NCT#) | | | | | |

Add any additional AEs reported for these drugs beyond the rows above.

### Table 3: Disappointing Drug Launches (Same Drug Class or Indications)

Find drugs in similar therapeutic areas that had BELOW-EXPECTATION commercial launches:

| Drug | Company | Indication | Approval Year | Analyst Consensus Peak ($M) | Actual Peak ($M) | Peak MS (%) | Reason for Underperformance | Company Size | Source |
|---|---|---|---|---|---|---|---|---|---|

### Table 4: Small Company Commercialization Challenges

Find small/mid-cap biotech companies (<$5B market cap at approval) that launched oncology drugs:

| Drug | Company | Market Cap at Launch ($B) | Indication | Year 1 Revenue ($M) | Year 3 Revenue ($M) | Salesforce Size | Co-Promotion Partner? | Key Challenge | Source |
|---|---|---|---|---|---|---|---|---|---|

### Table 5: Payer/Reimbursement Data

For the relevant indications, find data on drug pricing and reimbursement:

| Drug | Indication | WAC Price/Year ($K) | Net Price/Year ($K) | Medicare Coverage | Major PBM Coverage | Time to Formulary (months) | NCCN Guideline Category | Source |
|---|---|---|---|---|---|---|---|---|

---

IMPORTANT: This is DATA COLLECTION ONLY. Present raw data in tables. Do NOT predict, do NOT analyze, do NOT project market shares, do NOT give opinions. The analysis will be done separately by a human analyst.
"""

CATALYST_DATA_PROMPT = """
You are a senior biotech equity research DATA COLLECTOR. Your ONLY job is to COLLECT and ORGANIZE data.
Do NOT make predictions, do NOT project market shares, do NOT give opinions or analysis.
Just find and present ALL available data in structured tables.

## DRUG: {drug_name} ({company_name}, Ticker: {ticker})

## OBJECTIVE: Collect data needed for CATALYST SCENARIO analysis
Focus on: upcoming events calendar, historical data readout stock reactions,
comparable catalyst precedents, analyst consensus, and conference schedules.

## COLLECTED DATA (from Phase 1 comprehensive web search)
{data_report}

---

## YOUR TASK: Search for and present the following data tables.
Use Google Search extensively. Cite sources for EVERY number. If data is not publicly available, write "/".

### Table 1: Upcoming Catalyst Calendar

Find ALL upcoming events for {drug_name} and {company_name}:

| # | Event | Drug | Indication | Expected Date | Type (data/BLA/approval/conference) | Conference/Venue | NCT# | Trial Phase | Estimated Enrollment | Enrollment Status | Primary Endpoint | Source |
|---|---|---|---|---|---|---|---|---|---|---|---|---|

### Table 2: Historical Data Readouts + Stock Reactions (Peer View Format)

For EVERY past clinical data readout of {drug_name}, present in Peer View format:

| Field | Readout 1 | Readout 2 | Readout 3 | ... |
|---|---|---|---|---|
| Drug Name | | | | |
| Indication | | | | |
| NCT# | | | | |
| Treatment Line | | | | |
| Phase | | | | |
| Data Date | | | | |
| Conference/Venue | | | | |
| N (patients) | | | | |
| ORR (%) | | | | |
| BICR ORR (%) | | | | |
| CR (%) | | | | |
| PR (%) | | | | |
| DCR (%) | | | | |
| mPFS (months) | | | | |
| mOS (months) | | | | |
| 6 Mo PFS Rate (%) | | | | |
| 12 Mo PFS Rate (%) | | | | |
| ≥G3 AE Rate (%) | | | | |
| Stock Price Day Before ($) | | | | |
| Stock Price Day After ($) | | | | |
| Stock Change (%) | | | | |
| Trading Volume vs Avg (x) | | | | |

### Table 3: Comparable Catalyst Precedents (Same Drug Class/Indications)

For drugs with the SAME target, MOA, or indication as {drug_name}, find past data readout reactions:

| Field | Event 1 | Event 2 | Event 3 | Event 4 | ... |
|---|---|---|---|---|---|
| Drug Name | | | | | |
| Company | | | | | |
| Ticker | | | | | |
| Target/MOA | | | | | |
| Indication | | | | | |
| Treatment Line | | | | | |
| Phase | | | | | |
| Event Type (data/BLA/approval) | | | | | |
| Event Date | | | | | |
| Conference | | | | | |
| N (patients) | | | | | |
| ORR (%) | | | | | |
| CR (%) | | | | | |
| mPFS (months) | | | | | |
| mOS (months) | | | | | |
| ≥G3 AE Rate (%) | | | | | |
| Outcome (positive/negative/mixed) | | | | | |
| Stock Change −1d→+1d (%) | | | | | |
| Stock Change −1d→+5d (%) | | | | | |
| Market Cap at Event ($B) | | | | | |
| Prior Expectation (beat/miss/meet) | | | | | |

### Table 4: Analyst Consensus

Find analyst coverage and estimates for {company_name}/{drug_name}:

| Analyst | Firm | Date | Rating | Price Target ($) | Key Drug Assumption | Peak Revenue Estimate ($M) | Key Catalyst Mentioned | Source |
|---|---|---|---|---|---|---|---|---|

### Table 5: Upcoming Conference Calendar (Next 18 Months)

List all major oncology/biotech conferences where {company_name} may present:

| Conference | Dates | Location | Abstract Deadline | Presentation Type (oral/poster/N/A) | {company_name} Presenting? (confirmed/likely/unknown) | Source |
|---|---|---|---|---|---|---|

### Table 6: Standard-of-Care Benchmark Data

For each indication where a catalyst is expected, provide the current SOC efficacy benchmarks:

| Indication | Treatment Line | SOC Regimen | SOC ORR (%) | SOC CR (%) | SOC mPFS (months) | SOC mOS (months) | SOC ≥G3 AE Rate (%) | N | Source |
|---|---|---|---|---|---|---|---|---|---|

This defines the thresholds: beating SOC = positive catalyst, missing SOC = negative catalyst.

---

IMPORTANT: This is DATA COLLECTION ONLY. Present raw data in tables. Do NOT predict, do NOT analyze, do NOT project market shares, do NOT give opinions. The analysis will be done separately by a human analyst.
"""


# ══════════════════════════════════════════════════════════════════════════════
#  API CALLS
# ══════════════════════════════════════════════════════════════════════════════

def run_deep_research(client, prompt: str, drug_name: str) -> str:
    """Phase 1: Deep Research call for data collection.

    Returns the full data report as markdown text.
    """
    AGENT = "deep-research-pro-preview-12-2025"
    MAX_POLL = 3600
    POLL_INTERVAL = 30

    logger.info(f"  [Phase 1] Starting Deep Research data collection for {drug_name}...")

    interaction_id = None
    report = ""
    thought_count = 0
    stream_completed = False
    start_time = time.time()

    try:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)
            stream = client.interactions.create(
                agent=AGENT,
                input=prompt,
                background=True,
                store=True,
                stream=True,
            )

        for chunk in stream:
            elapsed = int(time.time() - start_time)

            if chunk.event_type == "interaction.start":
                interaction_id = chunk.interaction.id
                logger.info(f"  Interaction ID: {interaction_id}")

            elif chunk.event_type == "content.delta":
                if chunk.delta.type == "thought_summary":
                    thought_count += 1
                    first_line = chunk.delta.content.text.strip().split("\n")[0][:120]
                    if thought_count <= 5 or thought_count % 5 == 0:
                        logger.info(f"  [{elapsed}s] Thought #{thought_count}: {first_line}")
                elif chunk.delta.type == "text":
                    report += chunk.delta.text

            elif chunk.event_type == "interaction.complete":
                elapsed = int(time.time() - start_time)
                logger.info(f"  [{elapsed}s] Data collection completed via stream.")
                stream_completed = True
                break

    except Exception as stream_err:
        elapsed = int(time.time() - start_time)
        logger.info(f"  [{elapsed}s] Stream disconnected after {thought_count} thoughts. Polling...")

    if stream_completed and report:
        logger.info(f"  Data report: {len(report):,} chars, {thought_count} thoughts, {elapsed}s")
        return report

    # Polling fallback
    if not interaction_id:
        raise ValueError("Failed to get interaction ID")

    logger.info(f"  Polling interaction {interaction_id[:40]}...")
    poll_start = time.time()
    while (time.time() - poll_start) < MAX_POLL:
        time.sleep(POLL_INTERVAL)
        poll_elapsed = int(time.time() - poll_start)

        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)
            interaction = client.interactions.get(interaction_id)

        status = interaction.status
        if poll_elapsed % 120 == 0:
            logger.info(f"  [poll {poll_elapsed}s] Status: {status}")

        if status == "completed":
            text = interaction.outputs[-1].text
            if not text:
                raise ValueError("Empty output after polling")
            total = int(time.time() - start_time)
            logger.info(f"  [{total}s] Data collection completed via polling. {len(text):,} chars")
            return text
        elif status == "failed":
            raise RuntimeError(f"Deep Research failed: {getattr(interaction, 'error', 'unknown')}")

    raise TimeoutError("Deep Research data collection timed out")


def run_reasoning_pass(client, prompt: str, scenario: str, drug_name: str,
                       model: str = "gemini-2.5-flash") -> str:
    """Phase 2: Single reasoning pass using collected data as context.

    Uses streaming Interactions API to keep the connection alive during
    the model's long thinking phase. Falls back to generate_content on error.
    """
    MAX_RETRIES = 3

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            logger.info(f"  [Phase 2] Starting {scenario} reasoning for {drug_name} "
                        f"(attempt {attempt}/{MAX_RETRIES}, streaming)...")
            start_time = time.time()

            report = ""

            # Try interactions API with streaming first (keeps connection alive)
            try:
                with warnings.catch_warnings():
                    warnings.simplefilter("ignore", UserWarning)
                    stream = client.interactions.create(
                        model=model,
                        input=prompt,
                        stream=True,
                        store=True,
                        generation_config=genai_types.GenerateContentConfig(
                            temperature=0.7,
                            max_output_tokens=65535,
                        ),
                    )

                for chunk in stream:
                    elapsed = int(time.time() - start_time)
                    if chunk.event_type == "content.delta":
                        if hasattr(chunk.delta, 'type') and chunk.delta.type == "text":
                            report += chunk.delta.text
                        elif hasattr(chunk.delta, 'text') and chunk.delta.text:
                            report += chunk.delta.text
                    elif chunk.event_type == "interaction.complete":
                        break

                    # Progress logging every 60s
                    if elapsed > 0 and elapsed % 60 < 2 and report:
                        logger.info(f"    [{elapsed}s] {len(report):,} chars...")

            except Exception as stream_err:
                elapsed = int(time.time() - start_time)
                logger.info(f"    [{elapsed}s] Interactions stream error: {stream_err}")

                # Fallback: try models.generate_content_stream
                if not report:
                    logger.info(f"    Falling back to models.generate_content_stream...")
                    report = ""
                    gen_stream = client.models.generate_content_stream(
                        model=model,
                        contents=prompt,
                        config=genai_types.GenerateContentConfig(
                            temperature=0.7,
                            max_output_tokens=65535,
                        ),
                    )
                    for chunk in gen_stream:
                        if chunk.text:
                            report += chunk.text

            elapsed = int(time.time() - start_time)
            if not report:
                raise ValueError(f"Empty response for {scenario}")

            logger.info(f"  [{elapsed}s] {scenario} complete: {len(report):,} chars")
            return report

        except Exception as e:
            elapsed = int(time.time() - start_time) if 'start_time' in dir() else 0
            if attempt < MAX_RETRIES:
                wait = 30 * attempt
                logger.warning(f"  [{elapsed}s] Attempt {attempt} failed: {e}. "
                               f"Retrying in {wait}s...")
                time.sleep(wait)
            else:
                raise


# ══════════════════════════════════════════════════════════════════════════════
#  MARKDOWN → WORD (reuse from gemini_research.py)
# ══════════════════════════════════════════════════════════════════════════════

def _markdown_to_word(markdown_text: str, title: str) -> Document:
    """Convert markdown text to Word document with formatting."""
    doc = Document()
    h = doc.add_heading(title, level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ts = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    ts.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = ts.runs[0]
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph()

    lines = markdown_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            doc.add_paragraph()
            i += 1
            continue
        if line.startswith('##### '):
            doc.add_heading(line[6:], level=5)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                if not all(c.replace('-', '').replace(':', '').strip() == '' for c in row):
                    table_rows.append(row)
                i += 1
            if table_rows:
                max_cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=max_cols)
                table.style = 'Light Grid Accent 1'
                for ri, row_data in enumerate(table_rows):
                    for ci, cell_data in enumerate(row_data):
                        if ci < max_cols:
                            cell = table.rows[ri].cells[ci]
                            cell.text = cell_data
                            if ri == 0:
                                for r in cell.paragraphs[0].runs:
                                    r.font.bold = True
            continue
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(line, style='List Number')
        elif '**' in line:
            para = doc.add_paragraph()
            parts = line.split('**')
            for idx, part in enumerate(parts):
                r = para.add_run(part)
                if idx % 2 == 1:
                    r.font.bold = True
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                para = doc.add_paragraph('\n'.join(code_lines))
                para.style = 'No Spacing'
                if para.runs:
                    para.runs[0].font.name = 'Courier New'
                    para.runs[0].font.size = Pt(9)
            i += 1
            continue
        else:
            doc.add_paragraph(line)
        i += 1
    return doc


def save_report(report: str, output_dir: Path, ticker: str,
                drug_name: str, scenario: str) -> Path:
    """Save report as .md and .docx. Returns md_path."""
    output_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = re.sub(r'[^\w\-]', '_', drug_name)

    md_path = output_dir / f"{ticker}_{safe_name}_{scenario}_{ts}.md"
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(report)
    logger.info(f"  Saved: {md_path.name}")

    docx_path = output_dir / f"{ticker}_{safe_name}_{scenario}_{ts}.docx"
    title = f"{ticker} — {drug_name} {scenario.upper()} Analysis"
    doc = _markdown_to_word(report, title)
    doc.save(str(docx_path))
    logger.info(f"  Saved: {docx_path.name}")

    return md_path


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="Demo: one search → four independent thinking passes"
    )
    parser.add_argument("--ticker", required=True)
    parser.add_argument("--company-name", required=True)
    parser.add_argument("--drug", required=True, help="Drug name (e.g. CTX-009)")
    parser.add_argument("--trials-json", help="Path to clinical trials JSON")
    parser.add_argument("--data-report", help="Skip Phase 1: use existing data report .md file")
    parser.add_argument("--output-dir", help="Output directory")
    parser.add_argument("--model", default="gemini-2.5-flash",
                        help="Model for reasoning passes (default: gemini-2.5-flash)")
    parser.add_argument("--scenarios", nargs='+',
                        choices=["base", "bull", "bear", "catalyst"],
                        default=["base", "bull", "bear", "catalyst"],
                        help="Which scenarios to run (default: all four)")

    args = parser.parse_args()

    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        logger.error("GEMINI_API_KEY not set")
        sys.exit(1)

    import httpx
    # Long timeout: thinking models can take 2-5 min per reasoning pass
    http_options = genai_types.HttpOptions(timeout=600_000)  # 10 min
    client = genai.Client(api_key=api_key, http_options=http_options)

    output_dir = Path(args.output_dir) if args.output_dir else \
        Path(f"/mnt/c/Users/yzsun/Desktop/DD/{args.ticker}/pipeline_4pass_demo")

    # ── Load clinical trials data ──
    drug_trials_text = f"No pre-fetched trial data for {args.drug}. Search ClinicalTrials.gov."
    if args.trials_json:
        with open(args.trials_json) as f:
            trial_data = json.load(f)
        # Format for this drug
        from gemini_research import format_drug_trials
        drug_trials_text = format_drug_trials(args.drug, trial_data)

    # ══════════════════════════════════════════════════════════════════════
    #  PHASE 1: Data Collection (one Deep Research call)
    # ══════════════════════════════════════════════════════════════════════

    if args.data_report:
        logger.info(f"Skipping Phase 1 — loading data report from: {args.data_report}")
        with open(args.data_report, 'r', encoding='utf-8') as f:
            data_report = f.read()
        logger.info(f"  Data report: {len(data_report):,} chars")
    else:
        print(f"\n{'='*70}")
        print(f"PHASE 1: Deep Research Data Collection — {args.drug}")
        print(f"{'='*70}\n")

        data_prompt = DATA_COLLECTION_PROMPT.format(
            drug_name=args.drug,
            company_name=args.company_name,
            ticker=args.ticker,
            drug_trials_text=drug_trials_text,
        )

        data_report = run_deep_research(client, data_prompt, args.drug)

        # Save the data report
        data_path = save_report(data_report, output_dir, args.ticker, args.drug, "data")
        logger.info(f"  Data report saved: {data_path}")

    # ══════════════════════════════════════════════════════════════════════
    #  PHASE 2: Four Independent Reasoning Passes
    # ══════════════════════════════════════════════════════════════════════

    print(f"\n{'='*70}")
    print(f"PHASE 2: Independent Reasoning Passes — {args.drug}")
    print(f"  Model: {args.model}")
    print(f"  Scenarios: {', '.join(args.scenarios)}")
    print(f"  Data report: {len(data_report):,} chars")
    print(f"{'='*70}\n")

    scenario_prompts = {
        "base": BASE_REASONING_PROMPT,
        "bull": BULL_DATA_PROMPT,
        "bear": BEAR_DATA_PROMPT,
        "catalyst": CATALYST_DATA_PROMPT,
    }

    scenario_dirs = {
        "base": "pipeline_base4",
        "bull": "pipeline_bull2",
        "bear": "pipeline_bear3",
        "catalyst": "pipeline_catalyst",
    }

    results = []

    for scenario in args.scenarios:
        print(f"\n{'─'*60}")
        print(f"  {scenario.upper()} Case — {args.drug}")
        print(f"{'─'*60}")

        prompt = scenario_prompts[scenario].format(
            drug_name=args.drug,
            company_name=args.company_name,
            ticker=args.ticker,
            common_rules=_COMMON_RULES,
            data_report=data_report,
        )

        try:
            report = run_reasoning_pass(client, prompt, scenario, args.drug, args.model)

            # Save to scenario-specific directory
            scenario_dir = Path(f"/mnt/c/Users/yzsun/Desktop/DD/{args.ticker}/{scenario_dirs[scenario]}")
            md_path = save_report(report, scenario_dir, args.ticker, args.drug, "research")

            # Also save to demo output dir for comparison
            save_report(report, output_dir, args.ticker, args.drug, scenario)

            results.append({
                "scenario": scenario,
                "status": "SUCCESS",
                "chars": len(report),
                "path": str(md_path),
            })

        except Exception as e:
            logger.error(f"  FAILED {scenario}: {e}")
            import traceback
            traceback.print_exc()
            results.append({
                "scenario": scenario,
                "status": "FAILED",
                "error": str(e),
            })

        # Brief pause between reasoning calls
        if scenario != args.scenarios[-1]:
            time.sleep(5)

    # ══════════════════════════════════════════════════════════════════════
    #  SUMMARY
    # ══════════════════════════════════════════════════════════════════════

    print(f"\n{'='*70}")
    print(f"4-Pass Demo Complete — {args.ticker} / {args.drug}")
    print(f"{'='*70}")
    print(f"Data report: {len(data_report):,} chars")
    print(f"Output: {output_dir}\n")

    for r in results:
        if r["status"] == "SUCCESS":
            print(f"  ✓ {r['scenario']:10s} — {r['chars']:>8,} chars → {r['path']}")
        else:
            print(f"  ✗ {r['scenario']:10s} — FAILED: {r.get('error', 'unknown')}")

    succeeded = sum(1 for r in results if r["status"] == "SUCCESS")
    failed = sum(1 for r in results if r["status"] == "FAILED")
    print(f"\nTotal: {succeeded} succeeded, {failed} failed")
    print(f"{'='*70}\n")


if __name__ == "__main__":
    main()
