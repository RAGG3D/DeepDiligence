#!/usr/bin/env python3
"""
gemini_research.py -- Per-drug pipeline research via Gemini Deep Research Agent

Uses the Gemini Interactions API (Deep Research agent) for comprehensive,
multi-step web research on each pipeline asset.

Generates one detailed research report per pipeline asset, covering:
- Per-indication TAM, competitive landscape, differentiation
- Stage timeline predictions (Phase I→Approval)
- Post-launch market share forecasts (2024-2038) with data-backed reasoning

Usage:
    # Research all drugs for a ticker (auto-detect from ClinicalTrials.gov)
    python gemini_research.py --ticker CMPX --company-name "Compass Therapeutics"

    # Research specific drugs only
    python gemini_research.py --ticker CMPX --company-name "Compass Therapeutics" --drugs CTX-009 CTX-8371

    # Skip clinical trials fetching (use existing JSON)
    python gemini_research.py --ticker CMPX --company-name "Compass Therapeutics" --trials-json /path/to/trials.json
"""

import argparse
import json
import logging
import os
import re
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

try:
    from google import genai
    from google.genai import types as genai_types
except ImportError:
    print("ERROR: google-genai not installed. Run: pip install google-genai")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)


# ══════════════════════════════════════════════════════════════════════════════
#  CONSTANTS
# ══════════════════════════════════════════════════════════════════════════════

# Common standard-of-care drugs to filter out when identifying pipeline assets
_SOC_DRUGS = {
    "placebo", "pembrolizumab", "nivolumab", "atezolizumab", "durvalumab",
    "ipilimumab", "docetaxel", "paclitaxel", "carboplatin", "cisplatin",
    "gemcitabine", "pemetrexed", "bevacizumab", "trastuzumab", "rituximab",
    "cetuximab", "capecitabine", "fluorouracil", "5-fu", "oxaliplatin",
    "irinotecan", "doxorubicin", "cyclophosphamide", "etoposide",
    "temozolomide", "methotrexate", "vincristine", "prednisone",
    "dexamethasone", "lenalidomide", "bortezomib", "carfilzomib",
}


# ══════════════════════════════════════════════════════════════════════════════
#  PER-DRUG PROMPT TEMPLATE
# ══════════════════════════════════════════════════════════════════════════════

DRUG_RESEARCH_PROMPT = """
You are a senior biotech equity research analyst. Write a comprehensive, data-driven research report on the drug **{drug_name}** from **{company_name}** (Ticker: {ticker}).

This report has TWO parts:
1. **BASE CASE ANALYSIS** (Chapters 1-4): Full data collection + base case market share projections + stage timeline
2. **SUPPLEMENTARY DATA COLLECTION** (Chapters 5-8): Raw data tables for bull/bear/catalyst scenarios — NO predictions, NO opinions, just structured data

Use Google Search to find ALL relevant data. Every number MUST have a cited source.

---

## DRUG INFORMATION (from ClinicalTrials.gov)

{drug_trials_text}

---

## STRICT RULES (violating any rule invalidates the report)

**R1 — No vague indications**: NEVER write "Multiple indications", "Solid tumor(s)", "Advanced malignancies", "Various cancers". List EVERY specific cancer type by standard abbreviation (NSCLC, TNBC, BTC, CRC, mUC, HNSCC, OV, HL, Melanoma, etc.). Source each indication with NCT number or 10-K page.

**R2 — No unsupported numbers**: NEVER use "illustrative", "qualitative", "hypothetical". Every market share % and timeline year must be backed by cited data and explicit reasoning.

**R3 — Exhaustive competitors**: For each indication+line, list ALL marketed drugs AND ALL clinical-stage drugs targeting the SAME patient population. Estimate each clinical-stage competitor's approval year. You may only omit a clinical-stage competitor if it BOTH (a) shows clearly inferior data vs standard-of-care AND (b) is projected to launch 4+ years after this drug — and you must state the omission reason in one sentence.

**R4 — Line-matched comparison**: Compare 2L vs 2L, 3L vs 3L. NEVER compare 2L efficacy against 1L data. Match biomarker subgroups (PD-L1+ vs PD-L1+, EGFR+ vs EGFR+).

**R5 — Exhaust all reference comparables**: When predicting ANY number, list ALL relevant comparables (not just one convenient example). If 5 similar drugs exist, cite all 5. Apply the majority+proximity rule:
  - **Timeline predictions**: prioritize (1) same company size > (2) same molecule type > (3) same target+indication > (4) same mechanism
  - **Market share / peak share**: prioritize (1) same target+indication > (2) same mechanism+indication > (3) same molecule type > (4) same company size
  - **Ramp-up speed**: prioritize (1) same company size > (2) same target+indication > (3) same molecule type > (4) same mechanism

**R6 — TAM reference parameters** (from our DCF model — these take priority over web data):
  - Maturity Parameter: drugs reach peak share by Year 8 post-launch; 6% post-maturity growth; three tiers: Best-In-Class (Tagrisso curve) / Tier One (Alecensa curve) / Average
  - COGS benchmark: small molecule 37% (Xpovio), monoclonal antibody 45% (Danyelza)

**R7 — Exhaustive clinical data**: For EVERY data readout of this drug AND every competitor, include ALL officially published data points. Omitting any publicly available data point is a violation. If a data point is not publicly available, write "NR" (not reported).

---

## PART 1: BASE CASE ANALYSIS (Chapters 1-4)

### Chapter 1: Drug Overview
- Drug name, target, mechanism of action
- ALL known indications with source (NCT#, 10-K page, press release)
- Current clinical stage (highest phase across indications)
- Key differentiation vs existing therapies

### Chapter 2: Multi-Indication Forecasting Strategy
**MANDATORY PROCESS** — Before writing Chapter 3, you MUST:
1. List every indication and check: active trial? analyst discussion? company data disclosure?
2. Classify each indication: CLEAR prospects (separate forecast) vs UNCLEAR prospects (may combine)
3. State your decision: Separate All / Primary Separate + Secondary Combined / All Combined
4. Justify with: analyst coverage % per indication, social media sentiment, company guidance, TAM comparison
5. If combining, the combined name MUST list specific cancers (e.g. "HNSCC/NSCLC/OV Combined"), NEVER "Other Solid Tumors"

### Chapter 3: Per-Indication Analysis
**For EACH indication (or combined group), provide ALL of the following:**

#### 3.X.1 TAM (Total Addressable Market)
- Patient count for the SPECIFIC line+biomarker subgroup (not all-comers)
- 2024 patients, growth rate (CAGR), 2030 patients, 2038 patients
- Source: epidemiology reports, IQVIA, company investor presentation

#### 3.X.2 Comprehensive Efficacy Data — This Drug
**MANDATORY**: For EVERY clinical readout/data presentation of this drug in this indication, create a row. Include ALL officially published data — omitting any is a violation.

Table format (one row per readout event):
| Date | NCT# | Readout Phase | Treatment Line | Evaluable Patients | ORR | CR | PR | DCR | Median rPFS (mo) | Median PFS (mo) | Median OS (mo) | ≥G3 AEs (%) | ≥G3 SAE/Patients | Innovation/Result Summary | Market Reaction (-1d to +1d) | Source |

- **ORR/CR/PR/DCR**: Percentages. Include 95% CI if available: "35% (95% CI: 25-46%)"
- **≥G3 AEs**: Percentage of patients with Grade 3+ adverse events
- **Market Reaction**: Stock price change from -1 to +1 trading day around readout. "N/A" if not findable

#### 3.X.3 Competitive Landscape — Marketed Drugs
**MANDATORY**: Include EVERY marketed drug for this indication+line. EXHAUSTIVE — no omissions.

| Drug | Company | Line | Latest Sale (MM USD) | ORR | CR | PR | DCR | Median PFS (mo) | Median OS (mo) | ≥G3 AEs (%) | Route | Total Treatment Line | Median Treatment Line | Source |

#### 3.X.4 Competitive Landscape — Clinical-Stage Drugs

| Drug | Company | Phase | NCT# | Treatment Line | Evaluable Patients | ORR | CR | DCR | Median PFS (mo) | Median OS (mo) | ≥G3 AEs (%) | Est. Approval Year | Source |
- For each, explain how you estimated the approval year (cite comparables)
- If omitting a weak/late competitor, state: "Omitted [Drug]: [reason]"

#### 3.X.5 Differentiation Assessment
- Quantified comparison vs EACH competitor (ORR X% vs Y%, PFS X mo vs Y mo, CR X% vs Y%)
- Line-matched (2L vs 2L only)
- Safety comparison: ≥G3 AE rate for this drug vs each competitor
- Assessment: Best-in-class / Above-average / Average / Below-average
- Route advantage (oral vs IV)? Convenience advantage (dosing schedule)?

#### 3.X.6 Data Transfer Analysis (if applicable)
- Only if this indication has NO active trial but may be priced in
- Analyze: target expression overlap, pathology similarity, analyst mentions, 10-K listing
- Conclusion: priced in (lower peak, delayed launch) or not priced in (exclude)

#### 3.X.7 Market Share Projection — BASE Case (2024-2038)
| Year | Market Share | Reasoning |
| 2024 | 0% | Pre-approval |
| ... | ... | ... |
| 20XX | Peak X% | Peak based on [comparable: Drug A achieved Y% with similar ORR advantage] |
| 2038 | X% | ... |

**Peak share reasoning MUST include**: comparable drug(s), why higher/lower/equal, ALL comparables listed.

#### 3.X.8 Treatment Pricing Estimation
For EACH indication, estimate the treatment cost based on ALL marketed comparables in the same class:

| Parameter | Value | Source |
|-----------|-------|--------|
| Estimated Price Per Dose ($) | Based on all marketed same-class drugs | [cite sources] |
| Dosing Schedule | Q2W / Q3W / daily / etc. | [clinical protocol] |
| Treatment Duration (months) | Median PFS or standard cycle count | [cite trial data] |
| Doses Per Treatment Cycle | Number | [dosing protocol] |
| Total Treatment Cost Per Patient (MM USD) | Price x Doses x Duration | [calculation] |
| Route | IV / Oral / SC | [label/protocol] |

- List ALL marketed comparables used to derive the price estimate (drug, annual treatment cost, indication)
- If the drug is a novel modality (e.g. BsAb vs TKI), estimate pricing relative to BOTH the closest modality AND the indication benchmark
- For pre-approval drugs: estimate based on comparable approved drugs in the same target+indication

### Chapter 4: Stage Timeline
| Stage | Year | Source / Reasoning |
|-------|------|-------------------|
| 1 (Phase I) | XXXX | [Source: NCT/press release] |
| 2 (Phase II) | XXXX | [Source or prediction with ALL comparables listed] |
| 3 (Phase III) | XXXX | [Prediction: comparable drugs A(Xyr), B(Yyr), C(Zyr) -> median Wyr] |
| 4 (BLA Filing) | XXXX | [Prediction reasoning] |
| 5 (Approval) | XXXX | [Priority vs standard review reasoning] |

**Timeline comparables MUST include ALL similar drugs**: Drug name, same target/indication?, company size, Phase X->Y duration. Final prediction with majority+proximity justification.

---

## PART 2: SUPPLEMENTARY DATA COLLECTION (Chapters 5-7)

**IMPORTANT: Chapters 5-7 are DATA COLLECTION ONLY. Present raw data. Do NOT predict, do NOT analyze, do NOT project market shares, do NOT give opinions. Use "/" for data not publicly available.**

### Chapter 5: Peer View Data Collection

**CRITICAL**: For EACH indication where {drug_name} has a trial, exhaustively search for ALL marketed drugs AND ALL clinical-stage drugs (including private companies) targeting the SAME patient population. For EACH drug found, collect ALL historical clinical readouts.

**OUTPUT FORMAT**: Use the structured markers below. Each indication block is wrapped with `PEER_VIEW_START` / `PEER_VIEW_END`. Each drug readout is a simple key-value list. If a data point is not publicly available, write "/".

```
#### PEER_VIEW_START: {{INDICATION}}

##### Drug: {{DRUG_NAME}} — Readout 1
- Drug Name: {{name}}
- Company: {{company}}
- Ticker: {{ticker or "Private"}}
- Innovation: {{BIC/FIC/ADC/BsAb/Oral/etc.}}
- Target: {{target}}
- Result: {{Approved/Phase X/Continuing/Suspended}}
- NCT#: {{NCT number or /}}
- Treatment Line: {{1L/2L/3L+/adjuvant/neoadjuvant}}
- Phase: {{I/II/III}}
- Stage: {{1-5, where 1=Phase I, 5=Approved}}
- Data Date: {{YYYY-MM-DD}}
- Conference: {{conference name or journal or /}}
- N: {{evaluable patients}}
- ORR: {{percentage as decimal, e.g. 0.39}}
- BICR ORR: {{percentage as decimal or /}}
- CR: {{percentage as decimal or /}}
- PR: {{percentage as decimal or /}}
- DCR: {{percentage as decimal or /}}
- Median PFS: {{months or /}}
- Median rPFS: {{months or /}}
- Median OS: {{months or /}}
- 6 Mo PFS Rate: {{percentage as decimal or /}}
- 12 Mo PFS Rate: {{percentage as decimal or /}}
- 24 Mo PFS Rate: {{percentage as decimal or /}}
- 18 Mo OS Rate: {{percentage as decimal or /}}
- 24 Mo OS Rate: {{percentage as decimal or /}}
- Median DFS: {{months or /}}
- Median Follow-Up: {{months or /}}
- GEQ G3 SAE Pct: {{percentage as decimal or /}}
- GEQ G3 Clinical AE: {{most common Grade 3+ AEs or /}}
- Route: {{IV/Oral/SC}}
- Dosing Schedule: {{e.g. Q3W, Q2W, daily}}
- Latest Annual Sale: {{MM USD or /}}
- 1st Yr Sale: {{MM USD or /}}
- Stock Price Day Before: {{USD or /}}
- Stock Price Day After: {{USD or /}}
- Stock Change 1d: {{percentage as decimal or /}}
- Stock Change 3d: {{percentage as decimal or /}}
- Source: {{NCT#, publication, press release URL}}

##### Drug: {{DRUG_NAME}} — Readout 2
- Drug Name: ...
...

#### PEER_VIEW_END: {{INDICATION}}
```

**RULES for Chapter 5:**
- ALWAYS include {drug_name} as the FIRST entry in each indication block, even if no clinical data exists yet (use preclinical/IND data, write "/" for unavailable fields)
- Then marketed drugs (sorted by latest annual sale descending), then clinical-stage drugs (sorted by phase descending)
- For approved drugs: include the pivotal trial readout AND the most recent label-expansion readout
- For clinical-stage drugs: include EVERY publicly presented data readout
- Private companies: set Ticker to "Private"
- Stock prices: use the stock price on the trading day BEFORE and AFTER data presentation/publication. Write "/" if the company was private or data is not findable
- Each readout is ONE data presentation event (same drug may have multiple readouts)

### Chapter 6: Catalyst Data

#### 6.1 Upcoming Catalyst Calendar

| # | Event | Drug | Indication | Expected Date | Type | Conference/Venue | NCT# | Phase | Est. Enrollment | Enrollment Status | Primary Endpoint | Source |

#### 6.2 Historical Data Readouts + Stock Reactions

For EVERY past clinical data readout of {drug_name}, provide a quick-reference summary:

| # | Date | Indication | NCT# | Line | Phase | Conference | N | ORR | CR | mPFS | mOS | GEQ G3 AE | Stock -1d ($) | Stock +1d ($) | Change (%) | Source |

#### 6.3 Comparable Catalyst Precedents

Same target/MOA drugs — historical data disclosure events + stock reactions:

| # | Drug | Company | Ticker | Target/MOA | Indication | Line | Phase | Event Type | Date | Conference | N | ORR | CR | mPFS | mOS | GEQ G3 AE | Outcome | Stock Change 1d (%) | Stock Change 5d (%) | Mkt Cap ($B) | Source |

#### 6.4 Analyst Consensus

| Analyst | Firm | Date | Rating | PT ($) | Peak Rev Est ($M) | Key Catalyst | Source |

#### 6.5 Conference Calendar (Next 18 Months)

| Conference | Dates | Location | Abstract Deadline | {company_name} Presenting? | Source |

#### 6.6 Standard-of-Care Benchmark Data

For each indication where {drug_name} has a trial:

| Indication | Line | SOC Regimen | SOC ORR (%) | SOC CR (%) | SOC mPFS (mo) | SOC mOS (mo) | SOC GEQ G3 AE (%) | N | Source |

### Chapter 7: Source Summary
- All NCT numbers referenced
- All 10-K / press release citations
- All analyst report citations
- All epidemiology data sources
- All conference presentation citations

---

Begin comprehensive research now. Use Google Search extensively. Chapters 1-4: full base case analysis with predictions. Chapters 5-7: DATA COLLECTION ONLY — raw key-value data and tables, no opinions, no projections. Be exhaustive and cite all sources.
"""


# ══════════════════════════════════════════════════════════════════════════════
#  PER-CHAPTER PROMPT TEMPLATES
# ══════════════════════════════════════════════════════════════════════════════

PRICING_CHAPTER_PROMPT = """
You are a senior pharmaceutical pricing analyst.

Drug: {drug_name} ({company_name}, {ticker})
Indication(s): {indications}

## Drug Context
{drug_context}

## Task: Treatment Pricing Estimation (Chapter 3.X.8)

For EACH indication, research and estimate treatment pricing.
Use Google Search to find ALL comparable marketed drug prices.

Output format (one table per indication):

### [Indication Name]

| Parameter | Value | Source |
|-----------|-------|--------|
| Estimated Price Per Dose ($) | ... | ... |
| Dosing Schedule | Q2W/Q3W/daily/etc. | ... |
| Treatment Duration (months) | ... | ... |
| Doses Per Treatment Cycle | ... | ... |
| Annual Treatment Cost ($) | ... | ... |
| Total Treatment Cost Per Patient (MM USD) | ... | ... |
| Route | IV/Oral/SC | ... |

### Comparable Pricing Table
| Drug | Company | Indication | Annual Cost ($) | Per-Dose ($) | Route | Source |
|------|---------|------------|-----------------|--------------|-------|--------|
| ... | ... | ... | ... | ... | ... | ... |

RULES:
- List ALL marketed drugs in same indication+class
- Price must be US WAC or ASP (specify which)
- For pre-approval drugs: estimate based on class + indication positioning
- Total Cost = Price/dose x doses/cycle x cycles (= duration / cycle interval)
- Include both brand and biosimilar prices where applicable
"""


# ══════════════════════════════════════════════════════════════════════════════
#  CLINICAL TRIALS DATA FORMATTING
# ══════════════════════════════════════════════════════════════════════════════

def format_clinical_trials_data(trial_data: Dict) -> str:
    """Format ALL clinical trials data into readable text for Gemini prompt."""
    if not trial_data or not trial_data.get("trials"):
        return "No clinical trials data found."

    lines = []
    lines.append("### Company Clinical Trials (from ClinicalTrials.gov)\n")

    trials_by_drug: Dict[str, list] = {}
    for nct_id, trial in trial_data["trials"].items():
        for intervention in trial.get("interventions", []):
            if intervention not in trials_by_drug:
                trials_by_drug[intervention] = []
            trials_by_drug[intervention].append((nct_id, trial))

    for drug, trials in sorted(trials_by_drug.items()):
        if not drug or drug.lower() in _SOC_DRUGS:
            continue
        lines.append(f"**Drug: {drug}**\n")
        for nct_id, trial in trials:
            lines.append(f"- **{nct_id}**: {trial.get('title', 'N/A')}")
            lines.append(f"  - Phase: {trial.get('phase', 'Unknown')}")
            lines.append(f"  - Status: {trial.get('status', 'Unknown')}")
            lines.append(f"  - Conditions: {', '.join(trial.get('conditions', []))}")
            lines.append(f"  - Start: {trial.get('start_date', 'N/A')}")
            lines.append(f"  - Completion: {trial.get('completion_date', 'N/A')}\n")

    if trial_data.get("indications_summary"):
        lines.append("### Indications Summary\n")
        for ind, ncts in sorted(trial_data["indications_summary"].items(),
                                key=lambda x: len(x[1]), reverse=True):
            lines.append(f"- **{ind}**: {len(ncts)} trial(s) — {', '.join(ncts)}")

    return "\n".join(lines)


def format_drug_trials(drug_name: str, trial_data: Dict) -> str:
    """Format clinical trials data for a SINGLE drug."""
    if not trial_data or not trial_data.get("trials"):
        return f"No clinical trials data found for {drug_name}."

    lines = [f"### Clinical Trials for {drug_name}\n"]
    found = False

    for nct_id, trial in trial_data["trials"].items():
        interventions = [i.lower() for i in trial.get("interventions", [])]
        if drug_name.lower() in interventions or any(
            drug_name.lower() in i.lower() for i in trial.get("interventions", [])
        ):
            found = True
            lines.append(f"**{nct_id}**: {trial.get('title', 'N/A')}")
            lines.append(f"- Phase: {trial.get('phase', 'Unknown')}")
            lines.append(f"- Status: {trial.get('status', 'Unknown')}")
            lines.append(f"- Conditions: {', '.join(trial.get('conditions', []))}")
            lines.append(f"- Interventions: {', '.join(trial.get('interventions', []))}")
            lines.append(f"- Start: {trial.get('start_date', 'N/A')}")
            lines.append(f"- Completion: {trial.get('completion_date', 'N/A')}\n")

    if not found:
        lines.append(f"No specific trial records found for {drug_name} in the fetched data.")
        lines.append("Search ClinicalTrials.gov and company 10-K/press releases for trial details.")

    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
#  PIPELINE DRUG IDENTIFICATION
# ══════════════════════════════════════════════════════════════════════════════

def identify_pipeline_drugs(trial_data: Dict, ticker: str,
                            company_name: Optional[str] = None) -> List[Dict]:
    """
    Identify pipeline drugs from clinical trials data.

    Returns list of dicts: [{name, indications, phase, nct_ids}, ...]
    """
    if not trial_data or not trial_data.get("trials"):
        return []

    # Group by intervention
    drug_info: Dict[str, Dict] = {}
    for nct_id, trial in trial_data["trials"].items():
        for intervention in trial.get("interventions", []):
            name_lower = intervention.lower().strip()
            # Filter out SOC drugs, placebo, empty
            # Check both exact match and prefix match (e.g. "pembrolizumab (keytruda)")
            if not intervention or name_lower in _SOC_DRUGS or any(
                name_lower.startswith(soc) for soc in _SOC_DRUGS
            ):
                continue
            # Filter out generic descriptions
            if any(kw in name_lower for kw in [
                "standard", "best supportive", "investigator", "comparator"
            ]):
                continue

            if intervention not in drug_info:
                drug_info[intervention] = {
                    "name": intervention,
                    "indications": set(),
                    "phases": set(),
                    "nct_ids": [],
                }
            drug_info[intervention]["indications"].update(
                trial.get("conditions", [])
            )
            drug_info[intervention]["phases"].add(
                trial.get("phase", "Unknown")
            )
            drug_info[intervention]["nct_ids"].append(nct_id)

    # Convert sets to sorted lists
    result = []
    for drug in sorted(drug_info.values(), key=lambda d: d["name"]):
        result.append({
            "name": drug["name"],
            "indications": sorted(drug["indications"]),
            "phases": sorted(drug["phases"]),
            "nct_ids": sorted(drug["nct_ids"]),
        })

    return result


# ══════════════════════════════════════════════════════════════════════════════
#  GEMINI API
# ══════════════════════════════════════════════════════════════════════════════

def run_drug_research(ticker: str, company_name: str,
                      drug_name: str, drug_trials_text: str) -> str:
    """
    Run Gemini Deep Research for a SINGLE drug.

    Uses the Interactions API (deep-research-pro-preview-12-2025 agent)
    for comprehensive, multi-step web research.

    The unified prompt produces base/bull/bear + catalyst analysis in ONE call.
    Returns: Full research report as markdown text.
    """
    prompt = DRUG_RESEARCH_PROMPT.format(
        drug_name=drug_name,
        company_name=company_name,
        ticker=ticker,
        drug_trials_text=drug_trials_text,
    )

    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("GEMINI_API_KEY environment variable not set")

    client = genai.Client(api_key=api_key)

    logger.info(f"  Starting Deep Research for {drug_name} (hybrid stream+poll)...")

    DEEP_RESEARCH_AGENT = "deep-research-pro-preview-12-2025"
    MAX_POLL_TIME = 3600   # 60 min max for polling fallback
    POLL_INTERVAL = 30     # check every 30s

    import warnings

    last_exc = None
    for attempt in range(1, 4):
        try:
            # Phase 1: Start with streaming to get interaction ID + progress
            interaction_id = None
            report = ""
            thought_count = 0
            stream_completed = False

            try:
                with warnings.catch_warnings():
                    warnings.simplefilter("ignore", UserWarning)
                    stream = client.interactions.create(
                        agent=DEEP_RESEARCH_AGENT,
                        input=prompt,
                        background=True,
                        store=True,
                        stream=True,
                    )

                start_time = time.time()
                for chunk in stream:
                    elapsed = int(time.time() - start_time)

                    if chunk.event_type == "interaction.start":
                        interaction_id = chunk.interaction.id
                        logger.info(f"  Interaction ID: {interaction_id}")

                    elif chunk.event_type == "content.delta":
                        if chunk.delta.type == "thought_summary":
                            thought_count += 1
                            thought_text = chunk.delta.content.text.strip()
                            first_line = thought_text.split("\n")[0][:120]
                            if thought_count <= 5 or thought_count % 5 == 0:
                                logger.info(f"  [{elapsed}s] Thought #{thought_count}: {first_line}")
                        elif chunk.delta.type == "text":
                            report += chunk.delta.text

                    elif chunk.event_type == "interaction.complete":
                        elapsed = int(time.time() - start_time)
                        logger.info(f"  [{elapsed}s] Deep Research completed via stream.")
                        stream_completed = True
                        break

            except Exception as stream_err:
                elapsed = int(time.time() - start_time) if 'start_time' in dir() else 0
                logger.info(f"  [{elapsed}s] Stream disconnected after {thought_count} thoughts. "
                            f"Switching to polling...")

            # If stream gave us the full report, return it
            if stream_completed and report:
                logger.info(f"  {drug_name}: {len(report):,} chars, "
                            f"{thought_count} thoughts, {elapsed}s elapsed.")
                return report

            # Phase 2: Fall back to polling if stream broke
            if not interaction_id:
                raise ValueError("Failed to get interaction ID from stream")

            logger.info(f"  Polling interaction {interaction_id[:40]}...")
            poll_start = time.time()
            while (time.time() - poll_start) < MAX_POLL_TIME:
                time.sleep(POLL_INTERVAL)
                poll_elapsed = int(time.time() - poll_start)

                with warnings.catch_warnings():
                    warnings.simplefilter("ignore", UserWarning)
                    interaction = client.interactions.get(interaction_id)

                status = interaction.status
                if poll_elapsed % 120 == 0:
                    logger.info(f"  [poll {poll_elapsed}s] Status: {status}")

                if status == "completed":
                    text = interaction.outputs[-1].text
                    if not text:
                        raise ValueError(f"Empty output after polling for {drug_name}")
                    total = int(time.time() - start_time)
                    logger.info(f"  [{total}s] Deep Research completed via polling.")
                    logger.info(f"  {drug_name}: {len(text):,} chars generated.")
                    return text

                elif status == "failed":
                    error_msg = getattr(interaction, "error", "unknown error")
                    raise RuntimeError(f"Deep Research failed: {error_msg}")

            raise TimeoutError(
                f"Deep Research timed out after polling for {drug_name}"
            )

        except Exception as e:
            last_exc = e
            if attempt < 3:
                logger.warning(f"  Attempt {attempt} failed: {e}. Retrying in 15s...")
                time.sleep(15)
            else:
                raise last_exc

    raise last_exc  # Should never reach here


# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER-SPECIFIC RESEARCH (regular Gemini, not Deep Research)
# ══════════════════════════════════════════════════════════════════════════════

def run_chapter_research(ticker: str, company_name: str,
                         drug_name: str, chapter_prompt: str,
                         model: str = "gemini-2.5-flash") -> str:
    """
    Run regular Gemini Flash for a single chapter supplement.

    Uses Google Search grounding for up-to-date pricing data.
    NOT the Deep Research agent — faster, lower cost, focused scope.

    Returns: chapter text as markdown.
    """
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("GEMINI_API_KEY environment variable not set")

    client = genai.Client(api_key=api_key)

    logger.info(f"  Running chapter research for {drug_name} (model={model})...")

    response = client.models.generate_content(
        model=model,
        contents=chapter_prompt,
        config=genai_types.GenerateContentConfig(
            tools=[genai_types.Tool(google_search=genai_types.GoogleSearch())],
        ),
    )
    text = response.text
    logger.info(f"  {drug_name}: {len(text):,} chars generated")
    return text


def _extract_drug_context(report_dir: Path, ticker: str,
                          drug_name: str) -> str:
    """Extract drug context (Chapter 1 + 3.X.3 competitive data) from existing report."""
    safe_name = re.sub(r'[^\w\-]', '_', drug_name)
    pattern = f"{ticker}_{safe_name}_research_*.md"
    files = sorted(report_dir.glob(pattern), reverse=True)
    if not files:
        return f"No existing research report found for {drug_name}."

    content = files[0].read_text(encoding="utf-8")

    # Extract Chapter 1
    ch1_match = re.search(
        r'(?:^|\n)#+ *\*{0,2}\s*(?:Chapter 1|CHAPTER 1)[^\n]*\n(.*?)(?=\n#+ *\*{0,2}\s*(?:Chapter 2|CHAPTER 2)|\Z)',
        content, re.DOTALL | re.IGNORECASE
    )
    ch1_text = ch1_match.group(1)[:3000] if ch1_match else ""

    # Extract Chapter 3 competitive landscape sections (3.X.3)
    comp_sections = re.findall(
        r'(#{3,5}\s*3\.\d+\.3\s+[^\n]*\n.*?)(?=#{3,5}\s*3\.\d+\.\d+|\Z)',
        content, re.DOTALL
    )
    comp_text = "\n\n".join(s[:2000] for s in comp_sections[:5])

    parts = []
    if ch1_text:
        parts.append("### Chapter 1: Drug Overview\n" + ch1_text)
    if comp_text:
        parts.append("### Competitive Landscape Data\n" + comp_text)

    return "\n\n".join(parts) if parts else f"Report found but no extractable context for {drug_name}."


def _run_chapter_mode(args):
    """Handle --chapter mode: run single chapter research per drug."""
    output_dir = Path(args.output_dir) if args.output_dir else \
        Path(f"/mnt/c/Users/yzsun/Desktop/DD/{args.ticker}/pipeline_base4")
    report_dir = output_dir  # Existing reports are in the same directory

    if not args.drugs:
        logger.error("--chapter requires --drugs to specify which drugs to research")
        sys.exit(1)

    results = []
    for idx, drug_name in enumerate(args.drugs, 1):
        logger.info(f"\n{'─'*60}")
        logger.info(f"[{idx}/{len(args.drugs)}] Chapter '{args.chapter}' for: {drug_name}")
        logger.info(f"{'─'*60}")

        # Extract context from existing report
        drug_context = _extract_drug_context(report_dir, args.ticker, drug_name)

        # Build chapter prompt
        if args.chapter == "pricing":
            prompt = PRICING_CHAPTER_PROMPT.format(
                drug_name=drug_name,
                company_name=args.company_name,
                ticker=args.ticker,
                indications="(see drug context below)",
                drug_context=drug_context,
            )
        else:
            logger.error(f"Chapter '{args.chapter}' not yet implemented")
            continue

        try:
            text = run_chapter_research(
                args.ticker, args.company_name, drug_name, prompt)

            # Save chapter output
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_name = re.sub(r'[^\w\-]', '_', drug_name)
            md_path = output_dir / f"{args.ticker}_{safe_name}_{args.chapter}_{ts}.md"
            md_path.write_text(text, encoding="utf-8")
            logger.info(f"  Saved: {md_path}")

            results.append({"drug": drug_name, "status": "SUCCESS",
                           "chars": len(text), "file": str(md_path)})

            if idx < len(args.drugs):
                logger.info("  Pausing 5s before next drug...")
                time.sleep(5)

        except Exception as e:
            logger.error(f"  FAILED for {drug_name}: {e}")
            import traceback
            traceback.print_exc()
            results.append({"drug": drug_name, "status": "FAILED", "error": str(e)})

    # Summary
    print(f"\n{'='*70}")
    print(f"Chapter '{args.chapter}' Research Complete — {args.ticker}")
    print(f"{'='*70}")
    for r in results:
        if r["status"] == "SUCCESS":
            print(f"  OK  {r['drug']:20s} — {r['chars']:,} chars → {r['file']}")
        else:
            print(f"  ERR {r['drug']:20s} — {r.get('error', 'unknown')}")
    succeeded = sum(1 for r in results if r["status"] == "SUCCESS")
    print(f"\nTotal: {succeeded}/{len(results)} succeeded")
    print(f"{'='*70}\n")


# ══════════════════════════════════════════════════════════════════════════════
#  MARKDOWN → WORD CONVERSION
# ══════════════════════════════════════════════════════════════════════════════

def _markdown_to_word(markdown_text: str, title: str) -> Document:
    """Convert markdown text to Word document with formatting."""
    doc = Document()

    # Title
    h = doc.add_heading(title, level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    ts_para = doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    ts_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = ts_para.runs[0]
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph()

    lines = markdown_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            doc.add_paragraph()
            i += 1
            continue

        # Headers
        if line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)

        # Tables
        elif line.startswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                if not all(c.replace('-', '').replace(':', '').strip() == '' for c in row):
                    table_rows.append(row)
                i += 1
            if table_rows:
                max_cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=max_cols)
                table.style = 'Light Grid Accent 1'
                for ri, row_data in enumerate(table_rows):
                    for ci, cell_data in enumerate(row_data):
                        if ci < max_cols:
                            cell = table.rows[ri].cells[ci]
                            cell.text = cell_data
                            if ri == 0:
                                for run in cell.paragraphs[0].runs:
                                    run.font.bold = True
            continue

        # Bullet lists
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')

        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(line, style='List Number')

        # Bold text
        elif '**' in line:
            para = doc.add_paragraph()
            parts = line.split('**')
            for idx, part in enumerate(parts):
                r = para.add_run(part)
                if idx % 2 == 1:
                    r.font.bold = True

        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                para = doc.add_paragraph('\n'.join(code_lines))
                para.style = 'No Spacing'
                if para.runs:
                    para.runs[0].font.name = 'Courier New'
                    para.runs[0].font.size = Pt(9)
            i += 1
            continue

        # Regular text
        else:
            doc.add_paragraph(line)

        i += 1

    return doc


# ══════════════════════════════════════════════════════════════════════════════
#  SAVE REPORTS
# ══════════════════════════════════════════════════════════════════════════════

def save_drug_report(report: str, output_dir: Path, ticker: str,
                     drug_name: str) -> Tuple[Path, Path]:
    """Save per-drug report as .md and .docx. Returns (md_path, docx_path)."""
    output_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Sanitize drug name for filename
    safe_name = re.sub(r'[^\w\-]', '_', drug_name)

    md_path = output_dir / f"{ticker}_{safe_name}_research_{ts}.md"
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(report)
    logger.info(f"  Markdown: {md_path}")

    docx_path = output_dir / f"{ticker}_{safe_name}_research_{ts}.docx"
    title = f"{ticker} — {drug_name} Pipeline Research"
    doc = _markdown_to_word(report, title)
    doc.save(str(docx_path))
    logger.info(f"  Word doc: {docx_path}")

    return md_path, docx_path


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="Generate per-drug Gemini research reports for biotech pipeline"
    )
    parser.add_argument("--ticker", required=True, help="Stock ticker (e.g. CMPX)")
    parser.add_argument("--company-name", required=True, help="Full company name")
    parser.add_argument("--drugs", nargs='+',
                        help="Specific drug names to research (default: auto-detect all)")
    parser.add_argument("--trials-json", help="Path to existing clinical trials JSON")
    parser.add_argument("--output-dir",
                        help="Output directory (default: DD/{TICKER}/pipeline_base4/)")
    parser.add_argument("--skip-trials", action="store_true",
                        help="Skip ClinicalTrials.gov fetching")
    parser.add_argument("--chapter", choices=["pricing", "peer_view", "catalyst"],
                        help="Run single chapter research (supplements existing report)")

    args = parser.parse_args()

    # API key
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        logger.error("GEMINI_API_KEY not set")
        sys.exit(1)
    os.environ["GEMINI_API_KEY"] = api_key

    # Chapter mode: run single chapter per drug, then exit
    if args.chapter:
        _run_chapter_mode(args)
        return

    # Output directory
    if args.output_dir:
        output_dir = Path(args.output_dir)
    else:
        output_dir = Path(f"/mnt/c/Users/yzsun/Desktop/DD/{args.ticker}/pipeline_base4")

    # ── Step 1: Clinical Trials Data ──────────────────────────────────────────
    trial_data = None

    if args.trials_json:
        logger.info(f"Loading trials from {args.trials_json}")
        with open(args.trials_json) as f:
            trial_data = json.load(f)

    elif not args.skip_trials:
        try:
            from clinical_trials_fetcher import get_company_trials, save_trial_data
            logger.info("=" * 70)
            logger.info("STEP 1: Fetching Clinical Trials from ClinicalTrials.gov")
            logger.info("=" * 70)

            trial_data = get_company_trials(args.ticker, args.company_name)

            # Save JSON to output dir
            dd_dir = Path(f"/mnt/c/Users/yzsun/Desktop/DD/{args.ticker}")
            save_trial_data(args.ticker, trial_data, dd_dir)

            n_trials = len(trial_data.get("nct_ids", []))
            logger.info(f"Found {n_trials} clinical trials")
        except Exception as e:
            logger.warning(f"Failed to fetch clinical trials: {e}")

    # ── Step 2: Identify Pipeline Drugs ───────────────────────────────────────
    if args.drugs:
        # User specified drugs explicitly
        drugs = [{"name": d, "indications": [], "phases": [], "nct_ids": []}
                 for d in args.drugs]
        logger.info(f"User-specified drugs: {[d['name'] for d in drugs]}")
    elif trial_data:
        drugs = identify_pipeline_drugs(trial_data, args.ticker, args.company_name)
        logger.info(f"Auto-detected {len(drugs)} pipeline drugs: "
                    f"{[d['name'] for d in drugs]}")
    else:
        logger.error("No clinical trials data and no --drugs specified. "
                     "Use --drugs to specify drug names manually.")
        sys.exit(1)

    if not drugs:
        logger.error("No pipeline drugs found. Use --drugs to specify manually.")
        sys.exit(1)

    # ── Step 3: Per-Drug Gemini Research ──────────────────────────────────────
    logger.info("=" * 70)
    logger.info(f"STEP 2: Running Gemini Research for {len(drugs)} drug(s) [UNIFIED BASE+BULL+BEAR+CATALYST]")
    logger.info(f"Output: {output_dir}")
    logger.info("=" * 70)

    results = []
    for idx, drug in enumerate(drugs, 1):
        drug_name = drug["name"]
        logger.info(f"\n{'─'*60}")
        logger.info(f"[{idx}/{len(drugs)}] Researching: {drug_name}")
        logger.info(f"  Indications: {drug.get('indications', ['unknown'])}")
        logger.info(f"  Phases: {drug.get('phases', ['unknown'])}")
        logger.info(f"{'─'*60}")

        # Format drug-specific trials text
        if trial_data:
            drug_trials_text = format_drug_trials(drug_name, trial_data)
        else:
            drug_trials_text = (
                f"No pre-fetched trial data for {drug_name}. "
                "Search ClinicalTrials.gov and company filings for trial details."
            )

        try:
            report = run_drug_research(
                args.ticker, args.company_name,
                drug_name, drug_trials_text,
            )
            md_path, docx_path = save_drug_report(
                report, output_dir, args.ticker, drug_name,
            )
            results.append({
                "drug": drug_name,
                "status": "SUCCESS",
                "chars": len(report),
                "docx": str(docx_path),
            })

            # Pause between API calls to avoid rate limits
            if idx < len(drugs):
                logger.info("  Pausing 15s before next drug...")
                time.sleep(15)

        except Exception as e:
            logger.error(f"  FAILED for {drug_name}: {e}")
            import traceback
            traceback.print_exc()
            results.append({
                "drug": drug_name,
                "status": "FAILED",
                "error": str(e),
            })

    # ── Summary ───────────────────────────────────────────────────────────────
    print(f"\n{'=' * 70}")
    print(f"Pipeline Research Complete — {args.ticker} ({args.company_name}) [UNIFIED]")
    print(f"{'=' * 70}")
    print(f"Output directory: {output_dir}\n")

    for r in results:
        if r["status"] == "SUCCESS":
            print(f"  ✓ {r['drug']:20s} — {r['chars']:,} chars → {r['docx']}")
        else:
            print(f"  ✗ {r['drug']:20s} — FAILED: {r.get('error', 'unknown')}")

    succeeded = sum(1 for r in results if r["status"] == "SUCCESS")
    failed = sum(1 for r in results if r["status"] == "FAILED")
    print(f"\nTotal: {succeeded} succeeded, {failed} failed")
    print(f"{'=' * 70}\n")


if __name__ == "__main__":
    main()
